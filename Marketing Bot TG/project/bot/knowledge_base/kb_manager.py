"""
Управление базой знаний из PDF файлов
"""
import os
import logging
import time
import json
import sqlite3
from pathlib import Path
import PyPDF2
import docx  # Добавляем библиотеку для работы с DOCX
from bot.config import config
from bot.database import DBManager
import shutil
from docx import Document

logger = logging.getLogger(__name__)

class KnowledgeBaseManager:
    """
    Класс для управления базой знаний, включающей PDF документы.
    Обеспечивает загрузку, обработку, хранение и поиск по документам.
    """

    def __init__(self, db_manager=None):
        """
        Инициализация менеджера базы знаний

        Args:
            db_manager: Экземпляр DBManager для работы с базой данных
        """
        self.db_manager = db_manager or DBManager()
        self._setup_database()
        self.pdf_storage_path = Path(config.PDF_STORAGE_PATH)

        # Создаем директорию для хранения PDF, если она не существует
        os.makedirs(self.pdf_storage_path, exist_ok=True)

    def _setup_database(self):
        """Создать таблицы для базы знаний, если они не существуют"""
        # Таблица для PDF документов
        self.db_manager.execute_query('''
        CREATE TABLE IF NOT EXISTS knowledge_base_docs (
            doc_id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT UNIQUE,
            title TEXT,
            upload_date DATETIME,
            file_path TEXT,
            num_pages INTEGER,
            admin_id INTEGER
        )
        ''')

        # Таблица для контента PDF-файлов и индексации
        self.db_manager.execute_query('''
        CREATE TABLE IF NOT EXISTS knowledge_base_content (
            content_id INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_id INTEGER,
            page_num INTEGER,
            content TEXT,
            FOREIGN KEY (doc_id) REFERENCES knowledge_base_docs (doc_id)
        )
        ''')

        logger.info("База данных для базы знаний успешно подготовлена")

    def _extract_text_from_file(self, file_path):
        """
        Извлечь текст из PDF или DOCX файла

        Args:
            file_path: Путь к файлу

        Returns:
            dict: Словарь с текстом страниц {номер_страницы: текст}
            int: Количество страниц
        """
        pages_text = {}
        
        # Проверяем, является ли путь относительным
        if not os.path.isabs(file_path):
            # Если путь относительный, преобразуем его в абсолютный
            full_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), file_path)
            if os.path.exists(full_path):
                file_path = full_path
            else:
                logger.error(f"Файл не найден ни по относительному, ни по абсолютному пути: {file_path}, {full_path}")
                return {}, 0
        
        logger.info(f"Читаю файл по пути: {file_path}")
        
        # Проверяем существование файла
        if not os.path.exists(file_path):
            logger.error(f"Файл не существует: {file_path}")
            return {}, 0
            
        # Проверяем размер файла
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.error(f"Файл пустой (0 байт): {file_path}")
            return {}, 0
            
        logger.info(f"Размер файла: {file_size} байт")
        
        try:
            # Определяем тип файла по расширению
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                # Обработка PDF
                logger.info(f"Обрабатываю PDF файл: {file_path}")
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)
                    
                    logger.info(f"Количество страниц в PDF: {num_pages}")
                    
                    for page_num in range(num_pages):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        pages_text[page_num + 1] = text  # +1 для нумерации с 1
                        
                return pages_text, num_pages
            
            elif file_ext == '.docx':
                # Обработка DOCX
                logger.info(f"Обрабатываю DOCX файл: {file_path}")
                doc = docx.Document(file_path)
                
                # В DOCX нет страниц, поэтому мы будем делить текст на абзацы
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                num_paragraphs = len(paragraphs)
                
                logger.info(f"Количество абзацев в DOCX: {num_paragraphs}")
                
                # Группируем абзацы в "страницы" по 5 абзацев
                chunk_size = 5
                num_pages = (num_paragraphs + chunk_size - 1) // chunk_size  # округление вверх
                
                for page_num in range(num_pages):
                    start_idx = page_num * chunk_size
                    end_idx = min(start_idx + chunk_size, num_paragraphs)
                    text = "\n\n".join(paragraphs[start_idx:end_idx])
                    pages_text[page_num + 1] = text
                
                return pages_text, num_pages
            
            else:
                logger.error(f"Неподдерживаемый формат файла: {file_ext}")
                return {}, 0
                
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из файла {file_path}: {e}")
            return {}, 0

    # Переименовываем для обратной совместимости
    def _extract_text_from_pdf(self, pdf_path):
        """
        Извлечь текст из PDF файла (для обратной совместимости)

        Args:
            pdf_path: Путь к PDF файлу

        Returns:
            dict: Словарь с текстом страниц {номер_страницы: текст}
            int: Количество страниц
        """
        return self._extract_text_from_file(pdf_path)

    def load_document_directly(self, file_path, title=None):
        """
        Загрузить документ напрямую из локального пути (PDF или DOCX).
        
        Args:
            file_path: Путь к PDF или DOCX файлу
            title: Название документа (если None, будет использовано имя файла)
            
        Returns:
            dict: Результат операции
        """
        try:
            # Проверяем существование файла
            if not os.path.exists(file_path):
                error_msg = f"Файл не существует: {file_path}"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }
                
            # Проверяем размер файла
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                error_msg = f"Файл пустой (0 байт): {file_path}"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }
                
            # Определяем расширение файла
            ext = os.path.splitext(file_path)[1].lower()
            
            # Если название не указано, используем имя файла
            if not title:
                title = os.path.basename(file_path)
                
            # Обрабатываем в зависимости от типа файла
            if ext == '.pdf':
                return self.add_document_to_knowledge_base(file_path, title)
            elif ext == '.docx':
                return self._process_docx(file_path, title)
            else:
                error_msg = f"Неподдерживаемый формат файла: {ext}"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }
                
        except Exception as e:
            error_msg = f"Ошибка при загрузке документа: {str(e)}"
            logger.error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
            
    def _process_docx(self, file_path, title):
        """
        Обработать DOCX файл и добавить его в базу знаний
        
        Args:
            file_path: Путь к DOCX файлу
            title: Название документа
            
        Returns:
            dict: Результат операции
        """
        try:
            # Проверяем существование файла
            if not os.path.exists(file_path):
                error_msg = f"Файл не существует: {file_path}"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }
                
            # Проверяем размер файла
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                error_msg = f"Файл пустой (0 байт): {file_path}"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }
            
            filename = os.path.basename(file_path)
            
            # Создаем копию файла в хранилище
            target_path = os.path.join(self.pdf_storage_path, filename)
            
            # ВАЖНО: Проверяем, если целевой файл уже существует и не пустой,
            # мы НЕ перезаписываем его, а используем как есть
            if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                logger.info(f"Файл уже существует в целевой директории: {target_path}, используем его")
            
            # Копируем файл только если цель не существует, или если она пуста
            elif file_path != target_path:
                # Убеждаемся, что исходный файл не пустой перед копированием
                if file_size > 0:
                    logger.info(f"Копирую файл из {file_path} в {target_path}")
                    os.makedirs(os.path.dirname(target_path), exist_ok=True)
                    
                    # Используем shutil.copy2 вместо ручного копирования
                    shutil.copy2(file_path, target_path)
                    
                    # Проверяем, что копирование прошло успешно
                    if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                        logger.info(f"Файл успешно скопирован, размер: {os.path.getsize(target_path)} байт")
                    else:
                        error_msg = f"Ошибка при копировании: файл существует={os.path.exists(target_path)}, размер={os.path.getsize(target_path) if os.path.exists(target_path) else 'N/A'}"
                        logger.error(error_msg)
                        return {
                            'success': False,
                            'error': error_msg
                        }
                else:
                    error_msg = f"Исходный файл пустой (0 байт): {file_path}"
                    logger.error(error_msg)
                    return {
                        'success': False,
                        'error': error_msg
                    }
            
            # Извлекаем текст из DOCX файла
            try:
                doc = Document(target_path)
                content = ""
                
                for para in doc.paragraphs:
                    if para.text:
                        content += para.text + "\n"
                
                # Считаем количество абзацев как аналог страниц
                paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
                num_pages = max(1, paragraphs // 5)  # Примерно 5 абзацев на страницу
                
                if not content.strip():
                    error_msg = f"Не удалось извлечь текст из DOCX файла: {filename}"
                    logger.error(error_msg)
                    return {
                        'success': False,
                        'error': error_msg
                    }
                
            except Exception as e:
                error_msg = f"Ошибка при обработке DOCX файла: {str(e)}"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg
                }
            
            # Добавляем информацию о документе в БД
            current_time = time.strftime('%Y-%m-%d %H:%M:%S')
            result = self.db_manager.execute_query(
                """
                INSERT INTO knowledge_base_docs
                (filename, title, upload_date, file_path, num_pages, admin_id)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (filename, title, current_time, target_path, num_pages, 0)  # 0 = системная загрузка
            )

            if not result:
                return {
                    'success': False, 
                    'error': f"Ошибка при добавлении документа {filename} в базу данных"
                }

            # Получаем ID добавленного документа
            doc_id = self.db_manager.execute_query(
                "SELECT doc_id FROM knowledge_base_docs WHERE filename = ?",
                (filename,),
                fetch=True
            )

            if not doc_id or not doc_id[0]:
                return {
                    'success': False, 
                    'error': f"Не удалось получить ID для документа {filename}"
                }

            doc_id = doc_id[0][0]

            # Разбиваем текст на части (моделируем страницы)
            page_size = max(500, len(content) // num_pages)
            pages_text = {}
            
            for i in range(num_pages):
                start = i * page_size
                end = min((i + 1) * page_size, len(content))
                page_content = content[start:end]
                if page_content.strip():
                    pages_text[i + 1] = page_content
            
            # Добавляем контент каждой страницы в БД
            for page_num, content in pages_text.items():
                self.db_manager.execute_query(
                    """
                    INSERT INTO knowledge_base_content
                    (doc_id, page_num, content)
                    VALUES (?, ?, ?)
                    """,
                    (doc_id, page_num, content)
                )

            logger.info(f"DOCX файл {filename} успешно добавлен в базу знаний с ID {doc_id}")
            
            return {
                'success': True,
                'num_pages': num_pages,
                'doc_id': doc_id
            }
            
        except Exception as e:
            error_msg = f"Ошибка при добавлении DOCX в базу знаний: {str(e)}"
            logger.error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }

    def load_pdf_directly(self, file_path, title=None):
        """
        Метод для обратной совместимости, использует load_document_directly
        
        Args:
            file_path: Путь к PDF файлу
            title: Название документа
            
        Returns:
            tuple: (success, doc_id|error_message)
        """
        try:
            result = self.load_document_directly(file_path, title)
            
            if result['success']:
                return True, str(result['doc_id'])
            else:
                return False, result['error']
                
        except Exception as e:
            error_msg = f"Ошибка при загрузке PDF файла: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

    def add_pdf_to_knowledge_base(self, file_path, title, admin_id):
        """
        Добавить PDF файл в базу знаний

        Args:
            file_path: Путь к загруженному PDF файлу
            title: Название документа
            admin_id: ID администратора, который загрузил файл

        Returns:
            bool: Успешно ли добавлен документ
        """
        try:
            # Проверяем существование файла
            if not os.path.exists(file_path):
                logger.error(f"Файл {file_path} не существует")
                return False
                
            # Проверяем размер файла
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                logger.error(f"Файл пустой (0 байт): {file_path}")
                return False
                
            filename = os.path.basename(file_path)
            target_path = os.path.join(self.pdf_storage_path, filename)

            # ВАЖНО: Проверяем, если целевой файл уже существует и не пустой,
            # мы НЕ перезаписываем его, а используем как есть
            if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                logger.info(f"Файл уже существует в целевой директории: {target_path}, используем его")
            
            # Копируем файл только если цель не существует, или если она пуста
            elif file_path != target_path:
                # Убеждаемся, что исходный файл не пустой перед копированием
                if file_size > 0:
                    logger.info(f"Копирую файл из {file_path} в {target_path}")
                    os.makedirs(os.path.dirname(target_path), exist_ok=True)
                    
                    # Используем shutil.copy2 вместо ручного копирования
                    shutil.copy2(file_path, target_path)
                    
                    # Проверяем, что копирование прошло успешно
                    if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                        logger.info(f"Файл успешно скопирован, размер: {os.path.getsize(target_path)} байт")
                    else:
                        error_msg = f"Ошибка при копировании: файл существует={os.path.exists(target_path)}, размер={os.path.getsize(target_path) if os.path.exists(target_path) else 'N/A'}"
                        logger.error(error_msg)
                        return False
                else:
                    error_msg = f"Исходный файл пустой (0 байт): {file_path}"
                    logger.error(error_msg)
                    return False

            # Извлекаем текст из PDF
            pages_text, num_pages = self._extract_text_from_pdf(target_path)

            if num_pages == 0:
                logger.error(f"Не удалось извлечь текст из файла {filename}")
                return False

            # Добавляем информацию о документе в БД
            current_time = time.strftime('%Y-%m-%d %H:%M:%S')
            result = self.db_manager.execute_query(
                """
                INSERT INTO knowledge_base_docs
                (filename, title, upload_date, file_path, num_pages, admin_id)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (filename, title, current_time, target_path, num_pages, admin_id)
            )

            if not result:
                return False

            # Получаем ID добавленного документа
            doc_id = self.db_manager.execute_query(
                "SELECT doc_id FROM knowledge_base_docs WHERE filename = ?",
                (filename,),
                fetch=True
            )

            if not doc_id or not doc_id[0]:
                logger.error(f"Не удалось получить ID для документа {filename}")
                return False

            doc_id = doc_id[0][0]

            # Добавляем контент каждой страницы в БД
            for page_num, content in pages_text.items():
                self.db_manager.execute_query(
                    """
                    INSERT INTO knowledge_base_content
                    (doc_id, page_num, content)
                    VALUES (?, ?, ?)
                    """,
                    (doc_id, page_num, content)
                )

            logger.info(f"PDF файл {filename} успешно добавлен в базу знаний")
            return True

        except Exception as e:
            logger.error(f"Ошибка при добавлении PDF в базу знаний: {e}")
            return False

    def remove_pdf_from_knowledge_base(self, filename, admin_id):
        """
        Удалить PDF файл из базы знаний

        Args:
            filename: Имя файла для удаления
            admin_id: ID администратора, выполняющего удаление

        Returns:
            bool: Успешно ли удален документ
        """
        try:
            # Проверяем, существует ли документ и принадлежит ли он этому админу
            doc_info = self.db_manager.execute_query(
                "SELECT doc_id, file_path FROM knowledge_base_docs WHERE filename = ? AND admin_id = ?",
                (filename, admin_id),
                fetch=True
            )

            if not doc_info or not doc_info[0]:
                logger.warning(f"Попытка удаления несуществующего документа {filename} админом {admin_id}")
                return False

            doc_id, file_path = doc_info[0]

            # Удаляем контент документа
            self.db_manager.execute_query(
                "DELETE FROM knowledge_base_content WHERE doc_id = ?",
                (doc_id,)
            )

            # Удаляем запись о документе
            self.db_manager.execute_query(
                "DELETE FROM knowledge_base_docs WHERE doc_id = ?",
                (doc_id,)
            )

            # НЕ удаляем файл с диска
            if os.path.exists(file_path):
                logger.info(f"Файл {file_path} сохранен на диске для возможного повторного использования")

            logger.info(f"PDF файл {filename} успешно удален из базы знаний (без удаления файла с диска)")
            return True

        except Exception as e:
            logger.error(f"Ошибка при удалении PDF из базы знаний: {e}")
            return False

    def remove_pdf_by_id(self, doc_id):
        """
        Программный метод для удаления PDF файла по его ID без проверки прав администратора

        Args:
            doc_id: ID документа для удаления

        Returns:
            bool: Успешно ли удален документ
            str: Сообщение о результате операции
        """
        try:
            # Проверяем, существует ли документ
            doc_info = self.db_manager.execute_query(
                "SELECT filename, file_path FROM knowledge_base_docs WHERE doc_id = ?",
                (doc_id,),
                fetch=True
            )

            if not doc_info or not doc_info[0]:
                error_msg = f"Документ с ID {doc_id} не существует"
                logger.warning(error_msg)
                return False, error_msg

            filename, file_path = doc_info[0]

            # Удаляем контент документа
            self.db_manager.execute_query(
                "DELETE FROM knowledge_base_content WHERE doc_id = ?",
                (doc_id,)
            )

            # Удаляем запись о документе
            self.db_manager.execute_query(
                "DELETE FROM knowledge_base_docs WHERE doc_id = ?",
                (doc_id,)
            )

            # НЕ удаляем файл с диска
            if os.path.exists(file_path):
                logger.info(f"Файл {file_path} сохранен на диске для возможного повторного использования")

            success_msg = f"PDF файл {filename} (ID: {doc_id}) успешно удален из базы знаний (без удаления файла с диска)"
            logger.info(success_msg)
            return True, success_msg

        except Exception as e:
            error_msg = f"Ошибка при удалении PDF с ID {doc_id} из базы знаний: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

    def list_knowledge_base_docs(self, admin_id=None):
        """
        Получить список документов в базе знаний

        Args:
            admin_id: Опционально, ID админа для фильтрации документов

        Returns:
            list: Список документов в формате [{doc_id, filename, title, upload_date, num_pages}]
        """
        try:
            query = """
                SELECT doc_id, filename, title, upload_date, num_pages
                FROM knowledge_base_docs
            """
            params = ()

            if admin_id is not None:
                query += " WHERE admin_id = ?"
                params = (admin_id,)

            result = self.db_manager.execute_query(query, params, fetch=True)

            if not result:
                return []

            docs = []
            for row in result:
                docs.append({
                    'doc_id': row[0],
                    'filename': row[1],
                    'title': row[2],
                    'upload_date': row[3],
                    'num_pages': row[4]
                })

            return docs

        except Exception as e:
            logger.error(f"Ошибка при получении списка документов: {e}")
            return []

    def search_in_knowledge_base(self, query, limit=10):
        """
        Поиск в базе знаний по ключевым словам

        Args:
            query: Поисковый запрос
            limit: Максимальное количество результатов

        Returns:
            list: Список результатов поиска [{doc_id, title, page_num, snippet}]
        """
        try:
            # Простой поиск по вхождению фрагмента текста
            search_terms = query.lower().split()

            if not search_terms:
                return []

            # Подготавливаем запрос SQL с поиском по всем терминам
            sql_query = """
                SELECT d.doc_id, d.title, c.page_num, c.content
                FROM knowledge_base_content c
                JOIN knowledge_base_docs d ON c.doc_id = d.doc_id
                WHERE
            """

            conditions = []
            params = []

            for term in search_terms:
                conditions.append("LOWER(c.content) LIKE ?")
                params.append(f"%{term}%")

            sql_query += " AND ".join(conditions)
            sql_query += " LIMIT ?"
            params.append(limit)

            result = self.db_manager.execute_query(sql_query, tuple(params), fetch=True)

            if not result:
                return []

            search_results = []
            for row in result:
                doc_id, title, page_num, content = row

                # Создаем сниппет вокруг найденного текста
                snippet = self._create_snippet(content, search_terms[0], max_length=200)

                search_results.append({
                    'doc_id': doc_id,
                    'title': title,
                    'page_num': page_num,
                    'snippet': snippet
                })

            return search_results

        except Exception as e:
            logger.error(f"Ошибка при поиске в базе знаний: {e}")
            return []

    def _create_snippet(self, text, search_term, max_length=200):
        """
        Создать сниппет текста с контекстом вокруг найденного термина

        Args:
            text: Полный текст
            search_term: Искомый термин
            max_length: Максимальная длина сниппета

        Returns:
            str: Сниппет с контекстом
        """
        if not text:
            return ""

        lower_text = text.lower()
        lower_term = search_term.lower()
        pos = lower_text.find(lower_term)

        if pos == -1:
            # Если термин не найден, возвращаем начало текста
            return text[:max_length] + "..." if len(text) > max_length else text

        # Вычисляем начальную и конечную позиции сниппета
        half_length = max_length // 2
        start = max(0, pos - half_length)
        end = min(len(text), pos + len(search_term) + half_length)

        # Корректируем начало и конец, чтобы не разрывать слова
        if start > 0:
            while start > 0 and text[start] != ' ':
                start -= 1

        if end < len(text):
            while end < len(text) and text[end] != ' ':
                end += 1
                if end >= len(text) or end - start > max_length:
                    break

        snippet = text[start:end]

        # Добавляем маркеры начала и конца, если сниппет не совпадает с полным текстом
        if start > 0:
            snippet = "..." + snippet
        if end < len(text):
            snippet = snippet + "..."

        return snippet

    def get_content_for_query(self, query):
        """
        Получить релевантный контент из базы знаний для запроса пользователя

        Args:
            query: Запрос пользователя

        Returns:
            str: Объединенный релевантный контент для использования в запросе к GPT
        """
        search_results = self.search_in_knowledge_base(query, limit=5)

        if not search_results:
            return None

        # Собираем контент из найденных фрагментов
        relevant_content = []
        for result in search_results:
            doc_id = result['doc_id']
            page_num = result['page_num']
            title = result['title']

            # Получаем полный контент страницы
            page_content = self._get_page_content(doc_id, page_num)

            if page_content:
                relevant_content.append(f"Документ: {title} (Страница {page_num})\n\n{page_content}")

        return "\n\n".join(relevant_content)

    def _get_page_content(self, doc_id, page_num):
        """
        Получить полный контент страницы документа

        Args:
            doc_id: ID документа
            page_num: Номер страницы

        Returns:
            str: Текст страницы
        """
        try:
            result = self.db_manager.execute_query(
                "SELECT content FROM knowledge_base_content WHERE doc_id = ? AND page_num = ?",
                (doc_id, page_num),
                fetch=True
            )

            if result and result[0]:
                return result[0][0]
            return None

        except Exception as e:
            logger.error(f"Ошибка при получении контента страницы: {e}")
            return None

    def add_document_to_knowledge_base(self, file_path, title):
        """
        Добавить PDF файл в базу знаний (новый интерфейс)

        Args:
            file_path: Путь к загруженному PDF файлу
            title: Название документа

        Returns:
            dict: Результат операции с ключами:
                  - success: True/False
                  - error: Сообщение об ошибке (если есть)
                  - num_pages: Количество обработанных страниц
                  - doc_id: ID документа
        """
        try:
            # Проверяем существование файла
            if not os.path.exists(file_path):
                return {
                    'success': False, 
                    'error': f"Файл {file_path} не существует"
                }

            # Проверяем, что это PDF
            if not file_path.lower().endswith('.pdf'):
                return {
                    'success': False, 
                    'error': f"Файл {file_path} не является PDF файлом"
                }
                
            # Проверяем размер файла
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                return {
                    'success': False, 
                    'error': f"Файл пустой (0 байт): {file_path}"
                }

            filename = os.path.basename(file_path)
            
            # Создаем копию файла в хранилище
            target_path = os.path.join(self.pdf_storage_path, filename)

            # ВАЖНО: Проверяем, если целевой файл уже существует и не пустой,
            # мы НЕ перезаписываем его, а используем как есть
            if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                logger.info(f"Файл уже существует в целевой директории: {target_path}, используем его")
            
            # Копируем файл только если цель не существует, или если она пуста
            elif file_path != target_path:
                # Убеждаемся, что исходный файл не пустой перед копированием
                if file_size > 0:
                    logger.info(f"Копирую файл из {file_path} в {target_path}")
                    os.makedirs(os.path.dirname(target_path), exist_ok=True)
                    
                    # Используем shutil.copy2 вместо ручного копирования
                    shutil.copy2(file_path, target_path)
                    
                    # Проверяем, что копирование прошло успешно
                    if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                        logger.info(f"Файл успешно скопирован, размер: {os.path.getsize(target_path)} байт")
                    else:
                        error_msg = f"Ошибка при копировании: файл существует={os.path.exists(target_path)}, размер={os.path.getsize(target_path) if os.path.exists(target_path) else 'N/A'}"
                        logger.error(error_msg)
                        return {
                            'success': False,
                            'error': error_msg
                        }
                else:
                    error_msg = f"Исходный файл пустой (0 байт): {file_path}"
                    logger.error(error_msg)
                    return {
                        'success': False,
                        'error': error_msg
                    }

            # Извлекаем текст из PDF
            pages_text, num_pages = self._extract_text_from_pdf(target_path)

            if num_pages == 0:
                return {
                    'success': False, 
                    'error': f"Не удалось извлечь текст из файла {filename}"
                }

            # Добавляем информацию о документе в БД
            current_time = time.strftime('%Y-%m-%d %H:%M:%S')
            result = self.db_manager.execute_query(
                """
                INSERT INTO knowledge_base_docs
                (filename, title, upload_date, file_path, num_pages, admin_id)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (filename, title, current_time, target_path, num_pages, 0)  # 0 = системная загрузка
            )

            if not result:
                return {
                    'success': False, 
                    'error': f"Ошибка при добавлении документа {filename} в базу данных"
                }

            # Получаем ID добавленного документа
            doc_id = self.db_manager.execute_query(
                "SELECT doc_id FROM knowledge_base_docs WHERE filename = ?",
                (filename,),
                fetch=True
            )

            if not doc_id or not doc_id[0]:
                return {
                    'success': False, 
                    'error': f"Не удалось получить ID для документа {filename}"
                }

            doc_id = doc_id[0][0]

            # Добавляем контент каждой страницы в БД
            for page_num, content in pages_text.items():
                self.db_manager.execute_query(
                    """
                    INSERT INTO knowledge_base_content
                    (doc_id, page_num, content)
                    VALUES (?, ?, ?)
                    """,
                    (doc_id, page_num, content)
                )

            logger.info(f"PDF файл {filename} успешно добавлен в базу знаний с ID {doc_id}")
            
            return {
                'success': True,
                'num_pages': num_pages,
                'doc_id': doc_id
            }

        except Exception as e:
            error_msg = f"Ошибка при добавлении PDF в базу знаний: {str(e)}"
            logger.error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }

    def delete_document_from_knowledge_base(self, doc_id):
        """
        Удалить PDF файл из базы знаний по ID документа

        Args:
            doc_id: ID документа для удаления

        Returns:
            dict: Результат операции с ключами:
                  - success: True/False
                  - error: Сообщение об ошибке (если есть)
        """
        try:
            # Проверяем, существует ли документ
            doc_info = self.db_manager.execute_query(
                "SELECT file_path FROM knowledge_base_docs WHERE doc_id = ?",
                (doc_id,),
                fetch=True
            )

            if not doc_info or not doc_info[0]:
                return {
                    'success': False,
                    'error': f"Документ с ID {doc_id} не найден"
                }

            file_path = doc_info[0][0]

            # Удаляем контент документа
            self.db_manager.execute_query(
                "DELETE FROM knowledge_base_content WHERE doc_id = ?",
                (doc_id,)
            )

            # Удаляем запись о документе
            self.db_manager.execute_query(
                "DELETE FROM knowledge_base_docs WHERE doc_id = ?",
                (doc_id,)
            )

            # НЕ удаляем файл физически с диска, просто логируем
            if os.path.exists(file_path):
                logger.info(f"Файл {file_path} оставлен на диске для возможного повторного использования")

            logger.info(f"Документ с ID {doc_id} успешно удален из базы знаний, файл сохранен на диске")
            
            return {
                'success': True
            }

        except Exception as e:
            error_msg = f"Ошибка при удалении документа с ID {doc_id}: {str(e)}"
            logger.error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
